"""Resume text extraction - pypdf for PDF, python-docx for DOCX. No FastAPI
import here: plain functions over bytes, unit-testable without a request or a
Celery worker. Callers (the resume_parse task) turn ExtractionFailed's message
into candidates.parse_error.
"""

import io
import re

import pypdf
from docx import Document

MIN_CHARS = 50

# Icon fonts (phone/email/location glyphs, common in resume templates) map
# their pictograms into the Private Use Areas - pypdf faithfully returns
# whatever codepoint the font's cmap gives it, but there is no real character
# behind it. Left in, these glue onto adjacent words with no whitespace
# (observed on a real resume: an icon glyph directly against a phone number).
# Stripped to a space, not deleted outright, so words don't get glued
# together instead. Built from chr() rather than source escapes, to keep the
# exact codepoints unambiguous.
_PUA_RANGES = [(0xE000, 0xF8FF), (0xF0000, 0xFFFFD), (0x100000, 0x10FFFD)]
_PRIVATE_USE_AREAS = re.compile(
    "[" + "".join(f"{chr(lo)}-{chr(hi)}" for lo, hi in _PUA_RANGES) + "]"
)


def _strip_glyph_artifacts(text: str) -> str:
    return _PRIVATE_USE_AREAS.sub(" ", text)


class ExtractionFailed(Exception):
    """Message is the human-readable parse_error stored on the candidate."""


def _extract_pdf(content: bytes) -> str:
    try:
        reader = pypdf.PdfReader(io.BytesIO(content))
    except Exception as exc:
        raise ExtractionFailed(f"could not read PDF: {exc}") from exc
    if reader.is_encrypted:
        raise ExtractionFailed("PDF is password-protected")
    try:
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        raise ExtractionFailed(f"could not read PDF: {exc}") from exc


def _extract_docx(content: bytes) -> str:
    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise ExtractionFailed(f"could not read DOCX: {exc}") from exc
    # Resume templates commonly lay text out in tables (columns for skills,
    # contact info, etc.) rather than top-level paragraphs - both are read.
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def extract_text(content: bytes, ext: str) -> str:
    """`ext` comes from resume_validation.sniff_extension (magic bytes), never
    a filename - the same signal in local and cloud mode. Raises
    ExtractionFailed for anything unreadable; a result under MIN_CHARS is
    also a failure (scanned/image-only PDF), with a distinct message so it's
    never mistaken for a genuinely short resume."""
    if ext == "pdf":
        text = _extract_pdf(content)
    elif ext == "docx":
        text = _extract_docx(content)
    else:
        raise ExtractionFailed(f"unsupported file type for extraction: {ext}")

    text = _strip_glyph_artifacts(text)

    if len(text.strip()) < MIN_CHARS:
        raise ExtractionFailed(
            "no selectable text found (scanned or image-only document) - OCR is out of scope"
        )
    return text
